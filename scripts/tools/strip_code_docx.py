#!/usr/bin/env python3
"""strip_code_docx.py — strip NovaScript code blocks and output as Word (.docx).

Requires: python-docx (`pip install python-docx`)
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Optional

SCRIPT_PATH = Path(__file__).resolve()
PROJECT_ROOT = SCRIPT_PATH.parents[2]


class CliError(Exception):
    """An invocation or infrastructure error that returns exit code 2."""


DIALOGUE_PATTERN = re.compile(r'^(.+?)：：\u201c(.+?)\u201d\s*$')
CODE_BLOCK_START = re.compile(r'^@?<\|')
CODE_BLOCK_END = re.compile(r'\|>\s*$')


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Strip NovaScript code blocks and output as Word (.docx).",
    )
    parser.add_argument(
        "--input", "-i", metavar="PATH", required=True,
        help="input NovaScript scenario file",
    )
    parser.add_argument(
        "--output", "-o", metavar="PATH", required=True,
        help="output .docx file path",
    )
    parser.add_argument(
        "--keep-speaker-names", action="store_true",
        help="preserve speaker names before dialogue",
    )
    parser.add_argument(
        "--project", metavar="PATH",
        help="project root (defaults to repository root)",
    )
    return parser


def resolve_project_root(explicit: str | None) -> Path:
    if explicit:
        p = Path(explicit).resolve()
        if not p.is_dir():
            raise CliError(f"project root not found: {p}")
        return p
    return PROJECT_ROOT


def _parse_lines(filepath: Path) -> list[dict]:
    """Parse scenario into structured segments."""
    try:
        raw = filepath.read_text(encoding="utf-8")
    except Exception as exc:
        raise CliError(f"cannot read {filepath}: {exc}") from exc

    lines = raw.splitlines()
    segments: list[dict] = []
    in_code = False

    for line in lines:
        stripped = line.strip()

        if CODE_BLOCK_START.match(stripped):
            in_code = True
            continue
        if in_code:
            if CODE_BLOCK_END.match(stripped):
                in_code = False
            continue

        m = DIALOGUE_PATTERN.match(stripped)
        if m:
            segments.append({
                "type": "dialogue",
                "speaker": m.group(1).strip(),
                "text": m.group(2).strip(),
            })
        elif stripped:
            segments.append({"type": "narration", "text": stripped})

    return segments


def write_docx(filepath: Path, segments: list[dict], keep_speaker_names: bool, output: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise CliError(
            "python-docx is required. Install it with: pip install python-docx"
        )

    doc = Document()

    # Title
    title = doc.add_heading(filepath.stem, level=1)

    for seg in segments:
        if seg["type"] == "dialogue":
            if keep_speaker_names:
                p = doc.add_paragraph()
                run_speaker = p.add_run(seg["speaker"] + ": ")
                run_speaker.bold = True
                run_speaker.font.size = Pt(11)
                run_text = p.add_run(seg["text"])
                run_text.font.size = Pt(11)
            else:
                p = doc.add_paragraph(seg["text"])
                for run in p.runs:
                    run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(seg["text"])
            for run in p.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(80, 80, 80)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> None:
    args = _parser().parse_args()
    try:
        project_root = resolve_project_root(args.project)
    except CliError as e:
        print(str(e), file=sys.stderr)
        sys.exit(2)

    input_path = Path(args.input)
    if not input_path.is_absolute():
        cwd_fp = Path.cwd() / input_path
        proj_fp = project_root / input_path
        input_path = cwd_fp if cwd_fp.exists() else proj_fp
    if not input_path.is_file():
        print(f"strip_code_docx: error: input file not found: {input_path}", file=sys.stderr)
        sys.exit(2)

    output_path = Path(args.output)
    if not output_path.is_absolute():
        output_path = Path.cwd() / output_path

    try:
        segments = _parse_lines(input_path)
        write_docx(input_path, segments, args.keep_speaker_names, output_path)
        print(f"Written: {output_path} ({len(segments)} segments)")
    except CliError as e:
        print(f"strip_code_docx: error: {e}", file=sys.stderr)
        sys.exit(2)


if __name__ == "__main__":
    main()
